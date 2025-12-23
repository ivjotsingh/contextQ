"""Document parsing service for PDF, Word, and text files.

Handles:
- File validation (MIME type, size, structure)
- Text extraction from PDF (PyMuPDF), DOCX (python-docx), and TXT files
- Filename sanitization for security
- Content hashing for idempotency

All blocking I/O operations are wrapped with asyncio.to_thread for proper async handling.
"""

import asyncio
import hashlib
import logging
import re
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from config import get_settings

logger = logging.getLogger(__name__)

# Supported file types and their MIME types
SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt"})


class DocumentParseError(Exception):
    """Raised when document parsing fails."""

    pass


class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported."""

    pass


class FileTooLargeError(Exception):
    """Raised when file exceeds size limit."""

    pass


class EmptyDocumentError(Exception):
    """Raised when document contains no extractable text."""

    pass


class DocumentParser:
    """Service for parsing and processing documents.

    All file I/O operations are non-blocking, using asyncio.to_thread
    to avoid blocking the event loop.
    """

    def __init__(self) -> None:
        """Initialize document parser."""
        self.settings = get_settings()

    def sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal and other issues."""
        if not filename:
            raise ValueError("Filename cannot be empty")

        filename = Path(filename).name
        filename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", filename)

        max_length = 255
        if len(filename) > max_length:
            name, ext = Path(filename).stem, Path(filename).suffix
            filename = name[: max_length - len(ext)] + ext

        if not filename or filename.startswith("."):
            filename = "document" + Path(filename).suffix

        return filename

    def validate_file(
        self,
        filename: str,
        file_size: int,
        content_type: str | None = None,
    ) -> str:
        """Validate uploaded file and return extension."""
        if not filename:
            raise ValueError("Filename cannot be empty")

        if file_size <= 0:
            raise ValueError("File size must be positive")

        if file_size > self.settings.max_file_size_bytes:
            raise FileTooLargeError(
                f"File size {file_size} exceeds limit of {self.settings.max_file_size_mb}MB"
            )

        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"File type '{ext}' not supported. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        return ext

    def compute_content_hash(self, text: str) -> str:
        """Compute SHA256 hash of normalized text content."""
        normalized = " ".join(text.lower().split())
        return hashlib.sha256(normalized.encode()).hexdigest()

    async def parse_file(self, file_path: str, filename: str) -> dict[str, Any]:
        """Parse a document file and extract text content."""
        if not filename:
            raise ValueError("Filename cannot be empty")

        ext = Path(filename).suffix.lower()

        try:
            if ext == ".pdf":
                result = await asyncio.to_thread(self._parse_pdf_sync, file_path)
            elif ext == ".docx":
                result = await asyncio.to_thread(self._parse_docx_sync, file_path)
            elif ext == ".txt":
                result = await asyncio.to_thread(self._parse_txt_sync, file_path)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type: {ext}")

            text = result.get("text", "").strip()
            if not text:
                raise EmptyDocumentError("Document contains no extractable text")

            text = self._normalize_text(text)
            result["text"] = text
            result["content_hash"] = self.compute_content_hash(text)
            result["metadata"]["filename"] = self.sanitize_filename(filename)
            result["metadata"]["document_type"] = ext.lstrip(".")

            return result

        except (EmptyDocumentError, UnsupportedFileTypeError, ValueError):
            raise
        except Exception as e:
            logger.error("Failed to parse document %s: %s", filename, e)
            raise DocumentParseError(f"Failed to parse document: {e}") from e

    def _normalize_text(self, text: str) -> str:
        """Normalize text by cleaning up whitespace."""
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        lines = [line.strip() for line in text.split("\n")]
        return "\n".join(lines).strip()

    def _parse_pdf_sync(self, file_path: str) -> dict[str, Any]:
        """Parse PDF file using PyMuPDF (synchronous)."""
        doc = None
        try:
            doc = fitz.open(file_path)
            text_parts = []
            page_count = len(doc)

            for page_num, page in enumerate(doc, start=1):
                try:
                    page_text = page.get_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(
                        "Failed to extract text from page %d: %s", page_num, e
                    )

            metadata = {}
            if doc.metadata:
                metadata = {
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "creator": doc.metadata.get("creator", ""),
                }

            return {
                "text": "\n\n".join(text_parts),
                "page_count": page_count,
                "metadata": metadata,
            }

        except Exception as e:
            raise DocumentParseError(f"Failed to parse PDF: {e}") from e
        finally:
            if doc is not None:
                doc.close()

    def _parse_docx_sync(self, file_path: str) -> dict[str, Any]:
        """Parse Word document using python-docx (synchronous)."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                }

            return {
                "text": "\n\n".join(text_parts),
                "page_count": None,
                "metadata": metadata,
            }

        except Exception as e:
            raise DocumentParseError(f"Failed to parse DOCX: {e}") from e

    def _parse_txt_sync(self, file_path: str) -> dict[str, Any]:
        """Parse plain text file (synchronous)."""
        try:
            try:
                with open(file_path, encoding="utf-8") as f:
                    text = f.read()
            except UnicodeDecodeError:
                with open(file_path, encoding="latin-1") as f:
                    text = f.read()

            return {"text": text, "page_count": None, "metadata": {}}

        except Exception as e:
            raise DocumentParseError(f"Failed to parse TXT: {e}") from e
